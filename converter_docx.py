from docx import Document

import argparse
import re
import os
import magic


parser = argparse.ArgumentParser(
    description='Converts plain text files to docx and viceversa'
)

parser.add_argument(
    'filepath',
    metavar='path',
    type=str,
    nargs='+',
    help='File to convert'
)

parser.add_argument(
    "-v", "--verbose",
    help="increase output verbosity",
    action="store_true"
)

parser.add_argument(
    '-ext',
    nargs=1,
    default=['srt'],
    help='Extension to use on file out'
)

args = parser.parse_args()


def convert_all(file_paths):

    for path in file_paths:

        if not os.path.exists(path):
            print(f"File '{path}' doesn't exist, skipping...")
            continue
        
        if os.path.isdir(path):
            for file_path in [os.path.join(path, file_path) for file_path in os.listdir(path)]:
                if not os.path.isdir(file_path):
                    convert(file_path)
        else:
            convert(path)


def convert(file_in):
    basename = os.path.splitext(file_in)[0]
    file_type = magic.from_file(file_in)

    if file_type == "Microsoft OOXML":
        file_out = f"{basename}.{args.ext[0]}"
        docx_to_txt(file_in, file_out)
    else:
        try:
            file_out = f"{basename}.docx"
            txt_to_docx(file_in, file_out)
        except:
            print(f"Error on file {file_in}. {file_type} can't be converted into docx. Skipping...")


def txt_to_docx(file_in, file_out):

    if args.verbose:
        print(f"converting {file_in} to {file_out}")

    document = Document()

    file_in = open(file_in, "r", errors='replace').read()
    file_in = re.sub(r'[^\x00-\x7F]+|\x0c',' ', file_in)    # remove all non-XML-compatible characters

    document.add_paragraph(file_in)
    document.save(file_out)


def docx_to_txt(file_in, file_out):

    if args.verbose:
        print(f"converting {file_in} to {file_out}")

    file_o = open(file_out, "w")
    file_o.write(Document(file_in).paragraphs[0].text)
    file_o.close()


convert_all(args.filepath)