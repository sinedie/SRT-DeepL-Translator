from docx import Document

import argparse
import re
import os
import magic


parser = argparse.ArgumentParser(
    description='Converts txt to docx and viceversa.'
)

parser.add_argument(
    'filepath',
    metavar='path',
    type=str,
    nargs='+',
    help='Archivo a convertir'
)

args = parser.parse_args()

def convert_all(file_paths):

    for path in file_paths:

        if not os.path.exists(path):
            print(f"File '{path}' doesn't exist, skipping...")
            continue

        # paths = os.listdir(path) if os.path.isdir(path) else [path]

        # for file_path in paths:
        #     if not os.path.isdir(file_path):
        convert(path)


def convert(file_in):
    basename = os.path.splitext(file_in)[0]

    file_type = magic.from_file(file_in)

    if file_type == "Microsoft OOXML":
        file_out = f"{basename}.txt"
        docx_to_txt(file_in, file_out)
    elif file_type.startswith("ASCII text"):
        file_out = f"{basename}.docx"
        txt_to_docx(file_in, file_out)
    else:
        print(f"File type of {file_in} not recognized: {file_type}")


def txt_to_docx(file_in, file_out):
    document = Document()

    myfile = open(file_in).read()
    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters

    document.add_paragraph(myfile)
    document.save(file_out)


def docx_to_txt(fileIn, file_out):

    file_o = open(file_out, "w")
    file_o.write(Document(fileIn).paragraphs[0].text)
    file_o.close()

convert_all(args.filepath)